import os
import pickle
import numpy as np
from pathlib import Path
from python_speech_features import mfcc
from sklearn.utils import shuffle
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import wave
import scipy.io.wavfile as wav

# Завантаження моделі
pkl_knn = "pickle_model_mlp_MFCC0511.pkl"
with open(pkl_knn, 'rb') as file:
    pickle_model_mlp = pickle.load(file)

# Конфігурація
wlen = 0.5  # Довжина вікна (у секундах)
wstep = 1  # Крок вікна (у секундах)

# Функція для обробки одного файлу
def inference(input_file):
    try:
        rate, sig = wav.read(input_file)
        mfcc_feat = mfcc(sig, rate, winlen=wlen, winstep=wstep, nfilt=52, nfft=22050, numcep=52).astype(np.float32)
        
        y_prediction = pickle_model_mlp.predict(mfcc_feat)
        unique, counts = np.unique(y_prediction, return_counts=True)
        sum_pred = sum(counts).astype(float)
        result = dict(zip(unique, counts))
        
        prediction = max(result, key=result.get)
        percentage = (max(counts) / sum_pred) * 100

        with wave.open(str(input_file), 'rb') as f:  # Довжина файлу
            frames = f.getnframes()
            rate = f.getframerate()
            duration = frames / float(rate)

        return {
            "file_name": Path(input_file).name,
            "prediction": prediction,
            "percentage": '{:.2f}'.format(percentage),
            "duration": duration,
            "class_distribution": {cls: (count / sum_pred) * 100 for cls, count in result.items()}
        }
    except Exception as e:
        print(f"Помилка обробки {input_file}: {e}")
        return None

# Функція для аналізу папки
def analyze_folder(folder_path, output_word_path):
    document = Document()
    document.add_heading('Результати аналізу аудіофайлів', level=1)

    # Збір результатів
    results = []
    all_classes = set()

    for file in Path(folder_path).glob("*.wav"):
        result = inference(file)
        if result:
            results.append(result)
            all_classes.update(result["class_distribution"].keys())

    all_classes = sorted(all_classes)

    # Створення таблиці у Word
    table = document.add_table(rows=1, cols=len(all_classes) + 1, style='Light List Accent 1')
    header_row = table.rows[0].cells
    header_row[0].text = "Ім'я файлу"

    for i, cls in enumerate(all_classes):
        header_row[i + 1].text = cls

    for result in results:
        row_cells = table.add_row().cells
        row_cells[0].text = result["file_name"]
        for i, cls in enumerate(all_classes):
            row_cells[i + 1].text = f"{result['class_distribution'].get(cls, 0):.1f}%"

    # Побудова графіка
    plt.figure(figsize=(12, 6))
    for cls in all_classes:
        percentages = [result["class_distribution"].get(cls, 0) for result in results]
        plt.plot([result["file_name"] for result in results], percentages, label=cls, marker='o')

    plt.xlabel('Файли')
    plt.ylabel('Частка (%)')
    plt.title('Розподіл передбачених класів для кожного файлу')
    plt.legend(title="Класи")
    plt.xticks(rotation=45)
    plt.grid(axis='y', linestyle='--', alpha=0.7)
    plt.tight_layout()

    graph_path = "class_distribution.png"
    plt.savefig(graph_path)
    print(f"Графік збережено в {graph_path}")

    # Додати графік у документ
    document.add_heading('Графік розподілу класів', level=2)
    document.add_picture(graph_path, width=Inches(6))

    # Збереження документа
    document.save(output_word_path)
    print(f"Результати збережено в {output_word_path}")

# Використання функції
folder_path = "/Users/stepan_batih/ML/44KHZ/Silence/"
output_word_path = "audio_analysis.docx"
analyze_folder(folder_path, output_word_path)
