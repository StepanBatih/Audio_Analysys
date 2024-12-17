import os
import numpy as np
import pickle
import matplotlib.pyplot as plt
from tensorflow.keras.models import load_model
import scipy.io.wavfile as wav
from scipy.signal import resample
from docx import Document
from docx.shared import Inches

# Завантаження моделі, scaler і label encoder
model_path = "clear_data.h5"
scaler_path = "scaler_clear.pkl"
label_encoder_path = "label_encoder_clear.pkl"

model = load_model(model_path)

with open(scaler_path, 'rb') as file:
    scaler = pickle.load(file)

with open(label_encoder_path, 'rb') as file:
    encoder = pickle.load(file)

# Конфігурація
wlen_seconds = 0.025  # Довжина вікна (сек)
wstep_seconds = 0.05  # Крок між вікнами (сек)
expected_length = 200  # Очікувана кількість ознак (вхід для моделі)

# Функція для обробки аудіофайлу
def preprocess_audio(file_path):
    rate, sig = wav.read(file_path)

    # Якщо сигнал багатовимірний (наприклад, стерео), беремо лише один канал
    if sig.ndim > 1:
        sig = sig[:, 0]  # Вибираємо перший канал

    # Розрахунок кількості семплів на вікно
    window_length_samples = int(wlen_seconds * rate)
    window_step_samples = int(wstep_seconds * rate)

    # Розбиття сигналу на вікна
    raw_features = []
    for start in range(0, len(sig) - window_length_samples + 1, window_step_samples):
        window = sig[start:start + window_length_samples]
        # Масштабування кожного вікна до expected_length
        window_resampled = resample(window, expected_length)
        raw_features.append(window_resampled)

    raw_features = np.array(raw_features)

    # Перетворення в 2D масив
    raw_features = raw_features.reshape(raw_features.shape[0], -1)

    # Нормалізація
    raw_features = scaler.transform(raw_features)
    return raw_features

# Функція для аналізу файлів у папці
def analyze_folder(folder_path, output_word_path):
    document = Document()
    document.add_heading('Результати аналізу аудіофайлів', level=1)

    # Збір класів для заголовків таблиці
    all_classes = set()
    results = {}

    for file in os.listdir(folder_path):
        if file.endswith(".wav"):
            file_path = os.path.join(folder_path, file)
            try:
                X_test = preprocess_audio(file_path)
                predictions = model.predict(X_test)
                predicted_classes = np.argmax(predictions, axis=1)
                decoded_predictions = encoder.inverse_transform(predicted_classes)

                unique_classes, counts = np.unique(decoded_predictions, return_counts=True)
                percentages = (counts / len(decoded_predictions)) * 100
                results[file] = {cls: pct for cls, pct in zip(unique_classes, percentages)}

                all_classes.update(unique_classes)
            except Exception as e:
                print(f"Помилка обробки {file}: {e}")

    all_classes = sorted(all_classes)

    # Створення таблиці
    table = document.add_table(rows=1, cols=len(all_classes) + 1, style='Light List Accent 1')
    header_row = table.rows[0].cells
    header_row[0].text = "Ім'я файлу"

    for i, cls in enumerate(all_classes):
        header_row[i + 1].text = cls

    for file, class_distribution in results.items():
        row_cells = table.add_row().cells
        row_cells[0].text = file
        for i, cls in enumerate(all_classes):
            row_cells[i + 1].text = f"{class_distribution.get(cls, 0):.1f}%"

    # Створення графіка
    plt.figure(figsize=(12, 6))
    for cls in all_classes:
        percentages = [results[file].get(cls, 0) for file in results]
        plt.plot(list(results.keys()), percentages, label=cls, marker='o')

    plt.xlabel('Файли')
    plt.ylabel('Частка (%)')
    plt.title('Розподіл передбачених класів для кожного файлу')
    plt.legend(title="Класи")
    plt.xticks(rotation=45)
    plt.grid(axis='y', linestyle='--', alpha=0.7)
    plt.tight_layout()

    graph_path = "class_distribution.png"
    plt.savefig(graph_path)
    print(f"Графік збережено в {graph_path}")

    document.add_heading('Графік розподілу класів', level=2)
    document.add_picture(graph_path, width=Inches(6))

    # Збереження документа
    document.save(output_word_path)
    print(f"Результати збережено в {output_word_path}")

# Використання функції
folder_path = "/Users/stepan_batih/ML/44KHZ/Silence/"
output_word_path = "audio_analysis.docx"
analyze_folder(folder_path, output_word_path)
